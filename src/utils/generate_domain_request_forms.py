import argparse
import requests
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
from datetime import datetime
from dotenv import load_dotenv
import sys
import subprocess

def print_and_log(message):
    print(message)
    sys.stdout.flush()

def convert_to_pdf(docx_path, pdf_path):
    try:
        subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(pdf_path), docx_path], check=True)
    except subprocess.CalledProcessError as e:
        print_and_log(f"Failed to convert {docx_path} to PDF: {e}")
        sys.exit(1)

# Load environment variables from .env file
load_dotenv()
gov_lk_host = os.getenv('GOV_LK_HOST')

# Parse command-line arguments for the request token
parser = argparse.ArgumentParser(description='Generate a domain registration form.')
parser.add_argument('-t', '--token', required=True, help='Request token for the API')
args = parser.parse_args()
request_token = args.token

# Fetch data from the API
try:
    url = f"{gov_lk_host}/api/request/get-summary?requestToken={request_token}"
    response = requests.get(url)
    response.raise_for_status()  # Check for HTTP errors
    data = response.json()['data']
except requests.exceptions.RequestException as e:
    print_and_log(f"Failed to fetch data from API: {e}")
    sys.exit(1)

# Extract the necessary information
request_id = data["request_id"]
site_code = data["site_code"]
organization_name = data["organization_name"]
address = data["address"]
email = data["email"]
contact_no = data["contact_no"]
org_head_name = data["organization_head"]["full_name"]
org_head_designation = data["organization_head"]["designation"]
requested_domains = data["requested_domains"]
administrator_name = data["administrator"]["full_name"]
administrator_designation = data["administrator"]["designation"]
hosting_provider = data["hosting_provider"]
hosting_coordinator = data["hosting_coordinator"]
domains_list = ', '.join([domain['fqdn'] for domain in requested_domains])

project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '../'))
media_root = os.path.join(project_root, 'public/media/domain-request/forms/')
if not os.path.exists(media_root):
    try:
        os.makedirs(media_root)
    except Exception as e:
        print_and_log(f"Failed to create media directory: {e}")
        sys.exit(1)

# Create the Request Form Document
doc = Document()

# Set margins to narrow (0.5 inches from all sides)
sections = doc.sections
for section in sections:
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    # Add footer with request token and document version
    footer = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_text = footer.add_run(f"Request Token: {request_token} | Document Version: 2.1")
    footer_text.font.size = Pt(9)
    footer_text.font.color.rgb = RGBColor(0, 0, 0)

# Add centered title
title = doc.add_paragraph()
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = title.add_run(f"Domain Registration Form | Request ID - {request_id}")
run.font.name = 'Calibri'
run.font.size = Pt(14)
run.bold = True
title.add_run().add_break()
title.paragraph_format.line_spacing = None  # Ensure no line spacing
title.paragraph_format.space_after = Pt(0)  # Remove space after paragraph

# Add "Organization Information" heading
heading = doc.add_paragraph()
run = heading.add_run("Organization Information")
run.font.name = 'Calibri'
run.font.size = Pt(14)
run.bold = True
heading.paragraph_format.line_spacing = None  # Ensure no line spacing
heading.paragraph_format.space_after = Pt(0)  # Remove space after paragraph

# Add Organization Table with 6 rows (one for each data field) and 2 columns
org_table = doc.add_table(rows=6, cols=2)

# Define the data field names and corresponding data
field_names = ["Organization Name", "Address", "Email", "Contact No", "Organization Head", "Organization Head Designation"]
field_data = [
    organization_name,
    address,
    email,
    str(contact_no),
    org_head_name,
    org_head_designation
]

# Set table width to span the full width of the page minus the margins
table_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
col_widths = [table_width * 0.30, table_width * 0.70]

# Set column widths and populate the table with field names and data
for row_idx in range(6):
    row = org_table.rows[row_idx]
    for idx, width in enumerate(col_widths):
        cell = row.cells[idx]
        cell.width = width
        if idx == 0:
            cell.text = field_names[row_idx]
        else:
            cell.text = field_data[row_idx]
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.line_spacing = None
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)

# Function to set borders for a table
def set_table_borders(table):
    tbl = table._tbl  # Get the table xml element
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    
    border_values = {
        'top': {"val": "single", "sz": "4", "space": "0", "color": "auto"},
        'left': {"val": "single", "sz": "4", "space": "0", "color": "auto"},
        'bottom': {"val": "single", "sz": "4", "space": "0", "color": "auto"},
        'right': {"val": "single", "sz": "4", "space": "0", "color": "auto"},
        'insideH': {"val": "single", "sz": "4", "space": "0", "color": "auto"},
        'insideV': {"val": "single", "sz": "4", "space": "0", "color": "auto"},
    }

    for border_name, border_attrs in border_values.items():
        border = OxmlElement(f'w:{border_name}')
        for attr_name, attr_value in border_attrs.items():
            border.set(qn(f'w:{attr_name}'), attr_value)
        tblBorders.append(border)

    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.append(tblPr)

# Apply table borders
set_table_borders(org_table)

# Add a line break after the "Organization Information" table
line_break = doc.add_paragraph()
line_break.paragraph_format.line_spacing = None  # Ensure no line spacing
line_break.paragraph_format.space_after = Pt(0)  # Remove space after paragraph

# Add "Requesting Domain(s)" heading
domain_heading = doc.add_paragraph()
run = domain_heading.add_run("Requesting Domain(s)")
run.font.name = 'Calibri'
run.font.size = Pt(14)
run.bold = True
domain_heading.paragraph_format.line_spacing = None  # Ensure no line spacing
domain_heading.paragraph_format.space_after = Pt(0)  # Remove space after paragraph

# Add Requesting Domains Table with 2 columns: Domain and Reason
domain_table = doc.add_table(rows=1, cols=2)
domain_table.autofit = False

# Set header row
hdr_cells = domain_table.rows[0].cells
hdr_cells[0].text = 'Domain'
hdr_cells[1].text = 'Reason'
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.line_spacing = None
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.bold = True

# Populate the table with domain and reason
for domain in requested_domains:
    row_cells = domain_table.add_row().cells
    row_cells[0].text = domain['fqdn']
    row_cells[1].text = domain['reason']

    # Set column widths
    for row in domain_table.rows:
        row.cells[0].width = col_widths[0]
        row.cells[1].width = col_widths[1]

    for cell in row_cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.line_spacing = None  # Ensure default line spacing
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)

# Apply table borders
set_table_borders(domain_table)

# Add a line break after the "Requesting Domain(s)" table
line_break = doc.add_paragraph()
line_break.paragraph_format.line_spacing = None  # Ensure no line spacing
line_break.paragraph_format.space_after = Pt(0)  # Remove space after paragraph

# Add "Contact Information" heading
contact_heading = doc.add_paragraph()
run = contact_heading.add_run("Contact Information")
run.font.name = 'Calibri'
run.font.size = Pt(14)
run.bold = True
contact_heading.paragraph_format.line_spacing = None  # Ensure no line spacing
contact_heading.paragraph_format.space_after = Pt(0)  # Remove space after paragraph

# Function to create user table
def add_user_table(doc, role, user_data, hosting_provider=None):
    # Add role title
    role_title = doc.add_paragraph()
    role_title.keep_with_next = True  # Keep this paragraph on the same page as the next
    run = role_title.add_run(role)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.bold = True
    role_title.paragraph_format.line_spacing = None  # Ensure no line spacing
    role_title.paragraph_format.space_after = Pt(0)  # Remove space after paragraph
    
    # Create table
    user_table = doc.add_table(rows=0, cols=2)
    user_table.autofit = False
    
    # Add user data to the table, excluding 'id' and modifying parameter names
    user_data_mod = { 
        "Full Name": user_data.get("full_name"),
        "NIC No": user_data.get("nic"),
        "Mobile": user_data.get("mobile"),
        "Email": user_data.get("email"),
        "Designation": user_data.get("designation")
    }
    
    for key, value in user_data_mod.items():
        if value is not None:  # Only add rows for parameters that exist
            row_cells = user_table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = str(value)
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = None  # Ensure default line spacing
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
    
    if hosting_provider:
        row_cells = user_table.add_row().cells
        row_cells[0].text = "Hosting Place"
        row_cells[1].text = hosting_provider
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = None  # Ensure default line spacing
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
    
    # Set column widths
    for row in user_table.rows:
        row.cells[0].width = col_widths[0]
        row.cells[1].width = col_widths[1]
    
    # Apply table borders
    set_table_borders(user_table)
    
    # Add a line break after the table
    line_break = doc.add_paragraph()
    line_break.paragraph_format.line_spacing = None  # Ensure no line spacing
    line_break.paragraph_format.space_after = Pt(0)  # Remove space after paragraph

# Add user tables for each role
roles = ["Administrator Contact", "Technical Contact", "Content Developer", "Hosting Coordinator"]
users = [data["administrator"], data["technical_contact"], data["content_developer"], data["hosting_coordinator"]]
hostings = [None, None, None, hosting_provider]

for role, user, hosting in zip(roles, users, hostings):
    add_user_table(doc, role, user, hosting)

# Function to add DNS records tables
def add_dns_records(doc, domains):
    dns_heading = doc.add_paragraph()
    dns_heading.keep_with_next = True  # Keep this paragraph on the same page as the next
    run = dns_heading.add_run("DNS Records")
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.bold = True
    dns_heading.paragraph_format.line_spacing = None  # Ensure no line spacing
    dns_heading.paragraph_format.space_after = Pt(0)  # Remove space after paragraph

    for domain in domains:
        domain_title = doc.add_paragraph()
        domain_title.keep_with_next = True  # Keep this paragraph on the same page as the next
        run = domain_title.add_run(domain['fqdn'])
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.bold = True
        domain_title.paragraph_format.line_spacing = None  # Ensure no line spacing
        domain_title.paragraph_format.space_after = Pt(0)  # Remove space after paragraph

        dns_table = doc.add_table(rows=0, cols=3)
        dns_table.autofit = False

        records = {}
        for record in domain['dns_records']:
            record_type = record['type']
            if record_type not in records:
                records[record_type] = []
            record_details = []
            for k, v in record.items():
                if k not in ('dns_record_id', 'type_record_id', 'id', 'type'):
                    record_details.append((k, v))
            records[record_type].extend(record_details)

        for record_type, values in records.items():
            num_values = len(values)
            row_cells = dns_table.add_row().cells
            row_cells[0].text = f"{record_type} Record"
            row_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells[0].vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells[0].paragraphs[0].runs[0].bold = True

            for i in range(1, num_values):
                dns_table.add_row().cells

            for idx, (key, value) in enumerate(values):
                dns_table.rows[-num_values + idx].cells[1].text = key
                dns_table.rows[-num_values + idx].cells[2].text = str(value)
                for cell in dns_table.rows[-num_values + idx].cells[1:]:
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.line_spacing_rule = None  # Ensure default line spacing
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(0)
                        for run in paragraph.runs:
                            run.font.name = 'Calibri'
                            run.font.size = Pt(11)

            merged_cell = dns_table.cell(-num_values, 0)
            for i in range(1, num_values):
                merged_cell.merge(dns_table.cell(-num_values + i, 0))

        set_table_borders(dns_table)

# Add DNS Records section
add_dns_records(doc, requested_domains)

# Add a line break after the DNS Records section
line_break = doc.add_paragraph()
line_break.paragraph_format.line_spacing = None  # Ensure no line spacing
line_break.paragraph_format.space_after = Pt(0)  # Remove space after paragraph

# Add the "Confirmation Seal & Signature" title
confirmation_title = doc.add_paragraph()
run = confirmation_title.add_run("Confirmation Seal & Signature")
run.font.name = 'Calibri'
run.font.size = Pt(14)
run.bold = True
confirmation_title.paragraph_format.line_spacing = None  # Ensure no line spacing
confirmation_title.paragraph_format.space_after = Pt(0)  # Remove space after paragraph

# Add a line break before the confirmation text
line_break = doc.add_paragraph()
line_break.paragraph_format.line_spacing = None  # Ensure no line spacing
line_break.paragraph_format.space_after = Pt(0)  # Remove space after paragraph

# Add the confirmation text
confirmation_text = doc.add_paragraph()
confirmation_text.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
confirmation_text.keep_with_next = True  # Keep this paragraph on the same page as the next
run = confirmation_text.add_run(
    "I hereby confirm that the information provided above is true, correct, and accurate."
)
run.font.name = 'Calibri'
run.font.size = Pt(11)
run.bold = True
confirmation_text.paragraph_format.line_spacing = None  # Ensure no line spacing
confirmation_text.paragraph_format.space_after = Pt(0)  # Remove space after paragraph

# Add the combined administrator and organization head's name and designation table
combined_table = doc.add_table(rows=2, cols=2)
combined_table.autofit = False

# First row with blank lines
row_cells = combined_table.rows[0].cells
row_cells[0].text = '\n\n\n'  # 3 blank lines
row_cells[1].text = '\n\n\n'  # 3 blank lines

# Second row with centered names and designations
row_cells = combined_table.rows[1].cells
p = row_cells[0].add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run(administrator_name)
run.font.name = 'Calibri'
run.font.size = Pt(11)
run.bold = True

p = row_cells[0].add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run(administrator_designation)
run.font.name = 'Calibri'
run.font.size = Pt(11)

p = row_cells[1].add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run(org_head_name)
run.font.name = 'Calibri'
run.font.size = Pt(11)
run.bold = True

p = row_cells[1].add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run(org_head_designation)
run.font.name = 'Calibri'
run.font.size = Pt(11)

for row in combined_table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.line_spacing = None  # Ensure default line spacing
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)

set_table_borders(combined_table)

# Save the request form document
docx_request_filename = os.path.join(media_root, f"Request_Form_{site_code}_{request_id}.docx")
try:
    doc.save(docx_request_filename)
except Exception as e:
    print_and_log(f"Failed to save request form document: {e}")
    sys.exit(1)

# Convert the request form document to PDF
pdf_request_filename = os.path.join(media_root, f"Request_Form_{site_code}_{request_id}.pdf")
convert_to_pdf(docx_request_filename, pdf_request_filename)

# Delete the request form Word document after conversion
try:
    os.remove(docx_request_filename)
except Exception as e:
    print_and_log(f"Failed to delete request form document: {e}")
    sys.exit(1)

# Create the Cover Letter Document
cover_letter = Document()

# Set margins to narrow (0.5 inches from all sides)
sections = cover_letter.sections
for section in sections:
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    # Add footer with request token
    footer = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_text = footer.add_run(f"Request Token: {request_token}")
    footer_text.font.size = Pt(9)
    footer_text.font.color.rgb = RGBColor(0, 0, 0)

# Add cover letter content
paragraphs = [
    "",
    "",
    "",
    "",
    datetime.now().strftime("%Y-%m-%d"),
    "Hostmaster,",
    "Gov.lk Domain Registry,",
    "Network Operation Center,",
    "Information and Communication Technology Agency of Sri Lanka,",
    "490, R.A.DeMelMawatha,",
    "Colombo 03,",
    "Sri Lanka",
    "",
    "Request for the Domain Registration / Modification of domain(s) {},",
    "I hereby confirm that the information in the attached domain registration / modification form is accurate, "
    "and the request has been made for official purposes only. Please note that "
    "{} who is the {} of this organization will be the authorized officer "
    "for this request and you may contact him/her pertaining to this request in the future.",
    "Thank You.",
    "",
    "",
    "",
    "_______________________________",
    org_head_name,
    org_head_designation,
    organization_name
]

# Add paragraphs to the document
for para in paragraphs:
    p = cover_letter.add_paragraph()
    if para == "Request for the Domain Registration / Modification of domain(s) {},":
        run = p.add_run(para.format(domains_list))
        run.bold = True
    elif para.startswith("I hereby confirm that the information"):
        text_parts = para.split("{}")
        p.add_run(text_parts[0])
        run = p.add_run(administrator_name)
        run.bold = True
        p.add_run(text_parts[1])
        p.add_run(administrator_designation)
        p.add_run(text_parts[2])
    else:
        p.add_run(para)
    p.paragraph_format.line_spacing = None
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

# Save the cover letter document
docx_cover_letter_filename = os.path.join(media_root, f"Cover_Letter_{site_code}_{request_id}.docx")
try:
    cover_letter.save(docx_cover_letter_filename)
except Exception as e:
    print_and_log(f"Failed to save cover letter document: {e}")
    sys.exit(1)

# Convert the cover letter document to PDF
pdf_cover_letter_filename = os.path.join(media_root, f"Cover_Letter_{site_code}_{request_id}.pdf")
convert_to_pdf(docx_cover_letter_filename, pdf_cover_letter_filename)

# Delete the cover letter Word document after conversion
try:
    os.remove(docx_cover_letter_filename)
except Exception as e:
    print_and_log(f"Failed to delete cover letter document: {e}")
    sys.exit(1)

# Print the final paths of the generated PDF files
print(f"Request form generated: media/domain-request/forms/{os.path.basename(pdf_request_filename)}")
print(f"Cover letter generated: media/domain-request/forms/{os.path.basename(pdf_cover_letter_filename)}")
